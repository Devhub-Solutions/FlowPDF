"""
DOCX Parser Service v2
- Returns structured data: paragraphs separate from tables
- Tables returned as nested rows/cols for easy HTML rendering
- Mapping uses original_text + paragraph_index (no fragile offsets)
- Supports table loop (array) mappings for docxtpl {% for %}
"""

import os
import copy
import re
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.oxml.ns import qn


def parse_document_structure(file_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX file into a frontend-friendly structure.
    Returns paragraphs and tables as separate top-level arrays.
    Tables are nested: table → rows → cells.
    """
    doc = Document(file_path)

    # ─── Parse paragraphs ───
    paragraphs = []
    for para_idx, paragraph in enumerate(doc.paragraphs):
        runs_info = []
        for run_idx, run in enumerate(paragraph.runs):
            runs_info.append({
                "index": run_idx,
                "text": run.text,
                "bold": run.bold or False,
                "italic": run.italic or False,
                "underline": run.underline is not None and run.underline is not False,
                "font_name": run.font.name if run.font.name else None,
                "font_size": str(run.font.size) if run.font.size else None,
            })

        style_name = paragraph.style.name if paragraph.style else "Normal"
        paragraphs.append({
            "paragraph_index": para_idx,
            "text": paragraph.text,
            "runs": runs_info,
            "style": style_name,
            "is_heading": "Heading" in style_name,
            "alignment": str(paragraph.alignment) if paragraph.alignment else None,
        })

    # ─── Parse tables ───
    tables = []
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for col_idx, cell in enumerate(row.cells):
                cell_runs = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        cell_runs.append({
                            "text": run.text,
                            "bold": run.bold or False,
                            "italic": run.italic or False,
                        })
                cells.append({
                    "col_index": col_idx,
                    "text": cell.text,
                    "runs": cell_runs,
                })
            rows.append({
                "row_index": row_idx,
                "cells": cells,
            })

        tables.append({
            "table_index": table_idx,
            "num_rows": len(table.rows),
            "num_cols": len(table.columns) if table.rows else 0,
            "rows": rows,
        })

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "total_paragraphs": len(paragraphs),
        "total_tables": len(tables),
    }


# ═══════════════════════════════════════════
# TEXT REPLACEMENT (run-level precision)
# ═══════════════════════════════════════════

def _find_and_replace_in_paragraph(paragraph, target_text: str, replacement: str) -> bool:
    """
    Find target_text in a paragraph and replace it, preserving run formatting.
    Uses text-based matching instead of fragile offsets.
    Returns True if replacement was made.
    """
    full_text = paragraph.text
    if target_text not in full_text:
        return False

    start_offset = full_text.index(target_text)
    end_offset = start_offset + len(target_text)

    runs = paragraph.runs
    if not runs:
        return False

    # Build run position map
    run_positions = []
    current_pos = 0
    for idx, run in enumerate(runs):
        run_start = current_pos
        run_end = current_pos + len(run.text)
        run_positions.append((run_start, run_end, idx))
        current_pos = run_end

    # Find affected runs
    affected_runs = []
    for run_start, run_end, run_idx in run_positions:
        if run_end > start_offset and run_start < end_offset:
            affected_runs.append((run_start, run_end, run_idx))

    if not affected_runs:
        return False

    first_run_start, first_run_end, first_run_idx = affected_runs[0]
    last_run_start, last_run_end, last_run_idx = affected_runs[-1]

    prefix = runs[first_run_idx].text[:start_offset - first_run_start]
    suffix = runs[last_run_idx].text[end_offset - last_run_start:]

    if first_run_idx == last_run_idx:
        runs[first_run_idx].text = prefix + replacement + suffix
    else:
        runs[first_run_idx].text = prefix + replacement
        for _, _, run_idx in affected_runs[1:-1]:
            runs[run_idx].text = ""
        runs[last_run_idx].text = suffix

    return True


def _find_and_replace_in_cell(cell, target_text: str, replacement: str) -> bool:
    """Replace target_text in a table cell."""
    for para in cell.paragraphs:
        if _find_and_replace_in_paragraph(para, target_text, replacement):
            return True
    return False


# ═══════════════════════════════════════════
# APPLY MAPPINGS
# ═══════════════════════════════════════════

def apply_mappings_to_document(source_path: str, mappings: List[Dict], output_path: str) -> bool:
    """
    Apply all mappings to create a docxtpl-ready template.
    
    Mapping types:
      - "paragraph": replace original_text with {{label}} in a paragraph
      - "table_cell": replace original_text with {{label}} in a specific cell
      - "table_loop": mark a table row as a loop with {% for %} syntax
    """
    doc = Document(source_path)

    # ─── Separate mapping types ───
    para_mappings = [m for m in mappings if m.get("mapping_type", "paragraph") == "paragraph"]
    cell_mappings = [m for m in mappings if m.get("mapping_type") == "table_cell"]
    loop_mappings = [m for m in mappings if m.get("mapping_type") == "table_loop"]

    # ─── Apply paragraph mappings (text-based matching) ───
    for mapping in para_mappings:
        para_idx = mapping.get("paragraph_index", 0)
        original_text = mapping["original_text"]
        label = mapping["label"]
        jinja_var = "{{" + label + "}}"

        if para_idx < len(doc.paragraphs):
            _find_and_replace_in_paragraph(doc.paragraphs[para_idx], original_text, jinja_var)

    # ─── Apply table cell mappings ───
    for mapping in cell_mappings:
        table_idx = mapping.get("table_index", 0)
        row_idx = mapping.get("row_index", 0)
        col_idx = mapping.get("col_index", 0)
        original_text = mapping["original_text"]
        label = mapping["label"]

        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                jinja_var = "{{" + label + "}}"
                _find_and_replace_in_cell(cell, original_text, jinja_var)

    # ─── Apply table loop mappings ───
    # Instead of inserting {%tr} tags (which are unreliable),
    # we mark cells with __LOOP__ placeholders.
    # The actual row duplication happens in expand_table_loops() before docxtpl render.
    for mapping in loop_mappings:
        table_idx = mapping.get("table_index", 0)
        loop_var = mapping.get("loop_variable", "items")
        data_row_idx = mapping.get("data_row_index", 1)
        cell_labels = mapping.get("cell_labels", [])

        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            if data_row_idx < len(table.rows):
                data_row = table.rows[data_row_idx]

                # Replace cell contents with __LOOP__varname__fieldname__ markers
                for cell_map in cell_labels:
                    col = cell_map["col_index"]
                    label = cell_map["label"]
                    original = cell_map.get("original_text", "")

                    if col < len(data_row.cells):
                        cell = data_row.cells[col]
                        marker = f"__LOOP__{loop_var}__{label}__"
                        if original:
                            _find_and_replace_in_cell(cell, original, marker)
                        else:
                            for para in cell.paragraphs:
                                if para.runs:
                                    para.runs[0].text = marker
                                    for run in para.runs[1:]:
                                        run.text = ""

    doc.save(output_path)
    return True


def expand_table_loops(template_path: str, data: dict, mappings: list, output_path: str):
    """
    Manually expand table loop rows BEFORE docxtpl rendering.
    
    For each table_loop mapping:
    1. Find the template row in the table
    2. For each item in the array data, clone the row
    3. Replace __LOOP__varname__fieldname__ markers with actual values
    4. Remove the original template row
    
    This avoids docxtpl's unreliable {%tr} feature entirely.
    """
    import copy
    from lxml import etree
    doc = Document(template_path)
    
    loop_mappings = [m for m in mappings if m.get("mapping_type") == "table_loop"]
    
    for mapping in loop_mappings:
        table_idx = mapping.get("table_index", 0)
        loop_var = mapping.get("loop_variable", "items")
        data_row_idx = mapping.get("data_row_index", 1)
        cell_labels = mapping.get("cell_labels", [])
        
        items = data.get(loop_var, [])
        if not items or table_idx >= len(doc.tables):
            continue
            
        table = doc.tables[table_idx]
        if data_row_idx >= len(table.rows):
            continue
        
        # Get the template row XML element
        template_row = table.rows[data_row_idx]
        template_tr = template_row._tr
        parent = template_tr.getparent()
        
        # Clone and insert a new row for each item
        new_rows = []
        for item_idx, item in enumerate(items):
            if not isinstance(item, dict):
                item = {"value": str(item)}
                
            # Deep copy the template row
            new_tr = copy.deepcopy(template_tr)
            
            # Replace markers in all text elements
            for t_elem in new_tr.iter(qn('w:t')):
                if t_elem.text:
                    for cl in cell_labels:
                        marker = f"__LOOP__{loop_var}__{cl['label']}__"
                        if marker in t_elem.text:
                            value = str(item.get(cl['label'], ''))
                            t_elem.text = t_elem.text.replace(marker, value)
            
            new_rows.append(new_tr)
        
        # Insert all new rows after the template row
        for new_tr in reversed(new_rows):
            template_tr.addnext(new_tr)
        
        # Remove the original template row
        parent.remove(template_tr)
    
    doc.save(output_path)


